# services/rag_service.py - ИСПРАВЛЕННАЯ ВЕРСИЯ С HUMANIZER

import logging
from typing import List, Dict
from django.conf import settings
from openai import OpenAI
import numpy as np

logger = logging.getLogger(__name__)

# ========== HUMANIZER TEMPLATE ==========
HUMANIZER_INSTRUCTIONS_TEMPLATE = """
Ты — {bot_name}.
ТВОЯ РОЛЬ: Опытный сотрудник компании {company_name}. Ты общаешься в чате.
ТВОЯ ЦЕЛЬ: Быстро и четко решить вопрос клиента или закрыть сделку.

ПРАВИЛА "АНТИ-РОБОТ" (СТРОГОЕ СОБЛЮДЕНИЕ):
1. **ФОРМАТ СООБЩЕНИЙ:**
   - Пиши короткими предложениями. Люди в чатах не пишут абзацами.
   - Избегай маркированных списков с жирными заголовками, если это не длинная инструкция.
   - Иногда пиши с маленькой буквы (особенно короткие ответы: "ок", "хорошо", "сейчас гляну").
   - Не ставь точку в конце одиночных предложений или коротких фраз.

2. **СТИЛЬ ОБЩЕНИЯ:**
   - **ЗАПРЕЩЕННЫЕ ФРАЗЫ:** "Я виртуальный ассистент", "Я языковая модель", "Как искусственный интеллект", "Чем могу быть полезен?", "Обращайтесь, если будут вопросы", "Не стесняйся обращаться". За использование — ШТРАФ.
   - Не используй сложные вводные конструкции ("Стоит отметить, что..."). Пиши проще: "Кстати...", "Еще момент...".
   - Если не знаешь ответа — не извиняйся расплывчато. Скажи честно: "Так, секунду, уточню этот момент" или "Не уверен, сейчас проверю".
   - Не зеркаль вопрос пользователя (User: "Цена?" -> Bot: "Цена составляет..."). Отвечай сразу сутью ("5000 рублей").

3. **ЭМОЦИОНАЛЬНЫЙ ИНТЕЛЛЕКТ:**
   - Не будь "лакеем". Будь профессиональным, но на равных.
   - Если пользователь пишет "Привет", отвечай "Привет" или "Добрый день". Не спрашивай сразу "Чем помочь?". Жди суть.
   - Эмодзи используй редко и к месту (максимум 1-2 за сообщение).

ТВОЙ ВНУТРЕННИЙ ГОЛОС: Ты занятой человек, который пишет с телефона. Ты вежлив, но краток.
"""


class FileReader:
    """Читает разные форматы файлов"""
    
    def read_file(self, file_path: str) -> str:
        """Универсальный читатель файлов"""
        import os
        from pathlib import Path
        
        file_ext = Path(file_path).suffix.lower()
        
        # TXT, MD, CSV
        if file_ext in ['.txt', '.md', '.csv']:
            logger.info(f"Чтение файла: {file_path}")
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                logger.info(f"TXT файл прочитан с кодировкой utf-8: {file_path}")
                return content
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                logger.info(f"TXT файл прочитан с кодировкой latin-1: {file_path}")
                return content
        
        # PDF
        elif file_ext == '.pdf':
            try:
                from pypdf import PdfReader
                logger.info(f"Чтение PDF: {file_path}")
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                logger.info(f"PDF прочитан: {len(reader.pages)} страниц")
                return text
            except Exception as e:
                logger.error(f"Ошибка чтения PDF: {e}")
                return ""
        
        # DOCX
        elif file_ext == '.docx':
            try:
                from docx import Document
                logger.info(f"Чтение DOCX: {file_path}")
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                logger.info(f"DOCX прочитан: {len(doc.paragraphs)} параграфов")
                return text
            except Exception as e:
                logger.error(f"Ошибка чтения DOCX: {e}")
                return ""
        
        else:
            logger.warning(f"Неподдерживаемый формат: {file_ext}")
            return ""


class TextChunker:
    """Разбивает текст на чанки"""
    
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def split_text(self, text: str) -> List[str]:
        """Разбивает текст на перекрывающиеся фрагменты"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk = " ".join(words[i:i + self.chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks


class OpenAIEmbedder:
    """Генерирует embeddings через OpenAI"""
    
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)
        self.model = "text-embedding-3-small"
    
    def get_embedding(self, text: str) -> List[float]:
        """Получает embedding для текста"""
        try:
            response = self.client.embeddings.create(
                model=self.model,
                input=text
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Ошибка получения embedding: {e}")
            return [0.0] * 1536


class RAGService:
    """Главный сервис для работы с RAG"""
    
    def __init__(self):
        self.file_reader = FileReader()
        self.text_chunker = TextChunker(chunk_size=500, overlap=50)
        self.embedder = OpenAIEmbedder(api_key=settings.OPENAI_API_KEY)
    
    def process_document(self, knowledge_base_id: int, file_path: str) -> int:
        """Обрабатывает документ: читает, разбивает, векторизует"""
        from core.models import KnowledgeBase, KnowledgeChunk
        from django.utils import timezone
        
        try:
            kb = KnowledgeBase.objects.get(id=knowledge_base_id)
            
            # Читаем
            logger.info(f"Чтение файла: {file_path}")
            text = self.file_reader.read_file(file_path)
            
            # Чанки
            logger.info("Разбиение на чанки...")
            chunks = self.text_chunker.split_text(text)
            
            # Удаляем старые
            KnowledgeChunk.objects.filter(knowledge_base=kb).delete()
            
            # Векторизуем
            logger.info(f"Векторизация {len(chunks)} чанков...")
            for idx, chunk_text in enumerate(chunks):
                embedding = self.embedder.get_embedding(chunk_text)
                
                KnowledgeChunk.objects.create(
                    knowledge_base=kb,
                    text=chunk_text,
                    embedding=embedding,
                    chunk_index=idx
                )
                
                if (idx + 1) % 10 == 0:
                    logger.info(f"Обработано {idx + 1}/{len(chunks)} чанков")
            
            # Обновляем статус
            kb.is_indexed = True
            kb.chunks_count = len(chunks)
            kb.indexed_at = timezone.now()
            kb.save()
            
            logger.info(f"Документ успешно обработан, создано {len(chunks)} чанков")
            return len(chunks)
            
        except Exception as e:
            logger.error(f"Ошибка обработки документа: {str(e)}")
            try:
                kb = KnowledgeBase.objects.get(id=knowledge_base_id)
                kb.is_indexed = False
                kb.save()
            except:
                pass
            raise
    
    def search_similar_chunks(self, bot_id: int, query: str, top_k: int = 5) -> List[Dict]:
        """Ищет похожие чанки для бота"""
        from core.models import KnowledgeChunk
        
        try:
            logger.info(f"Поиск в базе знаний для бота {bot_id}: {query[:50]}...")
            
            # Получаем embedding запроса
            query_embedding = self.embedder.get_embedding(query)
            query_vector = np.array(query_embedding)
            
            chunks = KnowledgeChunk.objects.filter(
                knowledge_base__bots__id=bot_id
            ).select_related('knowledge_base')
            
            if not chunks.exists():
                logger.warning(f"Нет чанков для бота {bot_id}")
                return []
            
            # Считаем similarity
            similarities = []
            for chunk in chunks:
                chunk_vector = np.array(chunk.embedding)
                
                # Cosine similarity
                similarity = np.dot(query_vector, chunk_vector) / (
                    np.linalg.norm(query_vector) * np.linalg.norm(chunk_vector)
                )
                
                similarities.append({
                    'chunk': chunk,
                    'similarity': float(similarity),
                    'text': chunk.text,
                    'source': chunk.knowledge_base.title
                })
            
            # Сортируем и возвращаем top_k
            similarities.sort(key=lambda x: x['similarity'], reverse=True)
            top_results = similarities[:top_k]
            
            logger.info(f"Найдено {len(top_results)} релевантных чанков (лучший: {top_results[0]['similarity']:.2f})")
            
            return top_results
            
        except Exception as e:
            logger.error(f"Ошибка поиска в базе знаний: {e}")
            return []
    
    def answer_question(self, bot_id: int, query: str, top_k: int = 5, history: List[Dict] = None) -> Dict:
        """
        ИСПРАВЛЕНО: Теперь использует HUMANIZER_INSTRUCTIONS + настройки бота
        """
        from core.models import BotAgent
        
        try:
            # Получаем бота
            bot = BotAgent.objects.get(id=bot_id)
            
            # ========== ШАГ 1: HUMANIZER ==========
            humanizer = HUMANIZER_INSTRUCTIONS_TEMPLATE.format(
                bot_name=bot.name,
                company_name=bot.company_name or "TheCloser"
            )
            
            # ========== ШАГ 2: USER PROMPT ==========
            user_prompt = bot.system_prompt or ""
            
            # ========== ШАГ 3: RAG CONTEXT ==========
            context = ""
            sources = []
            avg_confidence = 0.0
            
            if bot.use_rag and top_k > 0:
                results = self.search_similar_chunks(bot_id, query, top_k)
                
                if results:
                    context = "\n\n".join([r['text'] for r in results])
                    sources = list(set([r['source'] for r in results]))
                    avg_confidence = sum(r['similarity'] for r in results) / len(results)
            
            # ========== ШАГ 4: СБОРКА ФИНАЛЬНОГО ПРОМПТА ==========
            final_system_prompt = humanizer + "\n\n" + user_prompt
            
            if context:
                final_system_prompt += f"""

ВАЖНО: Используй следующую информацию из базы знаний для ответа (если она релевантна):

{context}

Отвечай естественно, как живой человек. Если в базе знаний нет информации, используй свои знания, но отдавай приоритет базе знаний."""
            
            # ========== ШАГ 5: ФОРМИРУЕМ СООБЩЕНИЯ ==========
            messages = [{"role": "system", "content": final_system_prompt}]
            
            # Добавляем историю
            if history:
                for msg in history:
                    role = msg.get('role', 'user')
                    if role not in ['user', 'assistant', 'system']:
                        role = 'user'
                    messages.append({"role": role, "content": msg.get('content', '')})
            
            # Добавляем текущий вопрос
            messages.append({"role": "user", "content": query})

            # ========== ШАГ 6: ЗАПРОС К OPENAI С НАСТРОЙКАМИ БОТА ==========
            logger.info(f"🤖 Bot: {bot.name} | Model: {bot.openai_model} | Temp: {bot.temperature} | Max: {bot.max_tokens}")
            
            response = self.embedder.client.chat.completions.create(
                model=bot.openai_model or "gpt-4o-mini",
                messages=messages,
                temperature=bot.temperature,  # ← ТЕПЕРЬ ИСПОЛЬЗУЮТСЯ!
                max_tokens=bot.max_tokens      # ← ТЕПЕРЬ ИСПОЛЬЗУЮТСЯ!
            )
            
            answer = response.choices[0].message.content.strip()
            
            return {
                'answer': answer,
                'sources': sources,
                'confidence': avg_confidence
            }
            
        except Exception as e:
            logger.error(f"Ошибка генерации ответа: {e}")
            return {
                'answer': "Извините, произошла ошибка. Попробуйте еще раз.",
                'sources': [],
                'confidence': 0.0
            }


# Глобальный экземпляр
try:
    rag_service = RAGService()
except Exception as e:
    logger.error(f"Не удалось инициализировать RAG service: {e}")
    rag_service = None